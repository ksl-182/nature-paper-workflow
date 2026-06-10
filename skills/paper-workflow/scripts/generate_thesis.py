"""生成符合湖北工程学院本科毕业论文格式的 Word 文档
用法: python generate_thesis.py [content.json] [output.docx]
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import latex2mathml.converter
import re
import json
import sys
import os

# XSLT 样式表路径（MathML 转 OMML）
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_MML2OMML_XSL = os.path.join(_SCRIPT_DIR, 'MML2OMML.XSL')


def set_cell_border(cell, **kwargs):
    """设置单元格边框
    kwargs: top, bottom, left, right, insideH, insideV
    每个值为 dict: {"sz": "12", "val": "single", "color": "000000", "space": "0"}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, attrs in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        for key, val in attrs.items():
            element.set(qn(f'w:{key}'), val)
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_figure(doc, image_path, caption='', width=Inches(5)):
    """添加图片和图题
    image_path: 图片路径
    caption: 图题文字（如 "图1 系统架构"）
    width: 图片宽度
    """
    # 图片居中
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=width)

    # 图题
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(caption)
        set_font(run, '宋体', 'Times New Roman', Pt(10.5))

    return p


def add_section_break(doc, start_type='new_page'):
    """添加分节符"""
    from docx.enum.section import WD_SECTION_START
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    return new_section


def set_page_number(section, fmt='decimal', start=1):
    """设置页码格式
    fmt: 'decimal' 阿拉伯数字, 'roman' 罗马数字
    start: 起始页码
    """
    sectPr = section._sectPr
    # 页码格式
    pgNumType = OxmlElement('w:pgNumType')
    if fmt == 'roman':
        pgNumType.set(qn('w:fmt'), 'upperRoman')
    else:
        pgNumType.set(qn('w:fmt'), 'decimal')
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)


def add_page_number_field(paragraph, fmt='decimal'):
    """在段落中插入页码域"""
    # 添加 PAGE 域
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar2)

    # 设置字体
    for r in [run, run2, run3]:
        set_font(r, '宋体', 'Times New Roman', Pt(9))  # 六号字 ≈ 7.5pt


def add_toc(doc):
    """插入目录域"""
    # 目录标题
    p = add_paragraph_with_font(doc, '目  录', '黑体', 'Times New Roman',
                                 Pt(16), bold=True, align='center')
    doc.add_paragraph()

    # 插入 TOC 域
    p = doc.add_paragraph()

    # TOC 域代码
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fldChar2)

    # 占位文本
    run4 = p.add_run('（请右键此处，选择"更新域"以生成目录）')
    set_font(run4, '宋体', 'Times New Roman', Pt(12))

    run5 = p.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._element.append(fldChar3)

    doc.add_page_break()


def set_font(run, name_cn='宋体', name_en='Times New Roman', size=None, bold=False):
    """设置字体（中文+英文分别设置）"""
    run.font.name = name_en
    run.font.bold = bold
    if size:
        run.font.size = size
    # 设置中文字体
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name_cn)


def add_paragraph_with_font(doc, text, font_cn='宋体', font_en='Times New Roman',
                             size=Pt(12), bold=False, align=None, indent_first=None,
                             line_spacing=1.5):
    """添加段落并设置字体"""
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 首行缩进
    if indent_first:
        p.paragraph_format.first_line_indent = indent_first

    # 行距
    p.paragraph_format.line_spacing = line_spacing

    run = p.add_run(text)
    set_font(run, font_cn, font_en, size, bold)
    return p


def create_thesis_template():
    """创建湖北工程学院论文模板"""
    doc = Document()

    # ===== 页面设置 =====
    section = doc.sections[0]
    section.page_width = Cm(21.0)  # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    return doc


def add_cover_page(doc, title, college, major, student_id, name, advisor, date):
    """添加封面页"""
    # 空行调整位置
    for _ in range(6):
        doc.add_paragraph()

    # 学校名称
    p = add_paragraph_with_font(doc, '湖北工程学院', '黑体', 'Times New Roman',
                                 Pt(22), bold=True, align='center')

    # 论文类型
    p = add_paragraph_with_font(doc, '本科毕业论文（设计）', '黑体', 'Times New Roman',
                                 Pt(22), bold=True, align='center')
    doc.add_paragraph()

    # 题目
    p = add_paragraph_with_font(doc, title, '黑体', 'Times New Roman',
                                 Pt(16), bold=True, align='center')
    doc.add_paragraph()

    # 信息行
    info_items = [
        ('学    院：', college),
        ('专    业：', major),
        ('学    号：', student_id),
        ('姓    名：', name),
        ('指导教师：', advisor),
        ('完成日期：', date),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        set_font(run, '宋体', 'Times New Roman', Pt(14), bold=False)
        run = p.add_run(value)
        set_font(run, '宋体', 'Times New Roman', Pt(14), bold=False)

    # 分页
    doc.add_page_break()


def add_abstract_cn(doc, abstract, keywords):
    """添加中文摘要"""
    p = add_paragraph_with_font(doc, '摘  要', '黑体', 'Times New Roman',
                                 Pt(12), bold=True, align='center')
    doc.add_paragraph()

    # 摘要内容
    p = add_paragraph_with_font(doc, abstract, '宋体', 'Times New Roman',
                                 Pt(12), indent_first=Cm(0.74), line_spacing=1.5)

    # 关键词
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run('关键词：')
    set_font(run, '黑体', 'Times New Roman', Pt(12), bold=True)
    run = p.add_run(keywords)
    set_font(run, '宋体', 'Times New Roman', Pt(12))

    doc.add_page_break()


def add_abstract_en(doc, title_en, abstract_en, keywords_en):
    """添加英文摘要"""
    p = add_paragraph_with_font(doc, title_en, 'Times New Roman', 'Times New Roman',
                                 Pt(16), bold=True, align='center')
    doc.add_paragraph()

    p = add_paragraph_with_font(doc, 'Abstract', 'Times New Roman', 'Times New Roman',
                                 Pt(12), bold=True, align='left')
    doc.add_paragraph()

    # 摘要内容
    p = add_paragraph_with_font(doc, abstract_en, 'Times New Roman', 'Times New Roman',
                                 Pt(12), indent_first=Cm(0.74), line_spacing=1.5)

    # 关键词
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run('Key words: ')
    set_font(run, 'Times New Roman', 'Times New Roman', Pt(12), bold=True)
    run = p.add_run(keywords_en)
    set_font(run, 'Times New Roman', 'Times New Roman', Pt(12))

    doc.add_page_break()


def add_heading_1(doc, text):
    """一级标题：四号黑体加粗左对齐"""
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_font(run, '黑体', 'Times New Roman', Pt(14), bold=True)
    return p


def add_heading_2(doc, text):
    """二级标题：小四黑体加粗左对齐"""
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        set_font(run, '黑体', 'Times New Roman', Pt(12), bold=True)
    return p


def add_heading_3(doc, text):
    """三级标题：小四宋体加粗左对齐"""
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        set_font(run, '宋体', 'Times New Roman', Pt(12), bold=True)
    return p


def add_body(doc, text):
    """正文：小四宋体，1.5倍行距，首行缩进2字符"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)  # 2字符 ≈ 0.74cm

    # 处理加粗、斜体、行内代码
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[\d+\])', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, '宋体', 'Times New Roman', Pt(12), bold=True)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            set_font(run, '宋体', 'Times New Roman', Pt(12))
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_font(run, '宋体', 'Times New Roman', Pt(12))
        elif re.match(r'^\[\d+\]$', part):
            run = p.add_run(part)
            set_font(run, '宋体', 'Times New Roman', Pt(12))
            run.font.superscript = True
        else:
            if part:
                run = p.add_run(part)
                set_font(run, '宋体', 'Times New Roman', Pt(12))
    return p


def add_list_item(doc, text, ordered=False):
    """列表项"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_font(run, '宋体', 'Times New Roman', Pt(12))
    return p


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """设置单元格内边距（单位：twips, 1cm ≈ 567 twips）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def add_table_three_line(doc, headers, rows, table_num='', table_title=''):
    """三线表
    表头上下加粗线，表尾细线，无竖线
    """
    # 表序和表名
    if table_num:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(f'{table_num} {table_title}')
        set_font(run, '宋体', 'Times New Roman', Pt(12), bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 表格宽度占满页面
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    # 设置表格宽度为100%
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    # 设置表格无默认边框
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # 表头
    line = {"sz": "12", "val": "single", "color": "000000", "space": "0"}
    none = {"val": "none", "sz": "0", "space": "0"}

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_font(run, '宋体', 'Times New Roman', Pt(12), bold=True)
        set_cell_border(cell, top=line, bottom=line, left=none, right=none)
        set_cell_margins(cell, top=40, bottom=40, start=80, end=80)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        is_last = (row_idx == len(rows) - 1)
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            set_font(run, '宋体', 'Times New Roman', Pt(12))
            set_cell_margins(cell, top=40, bottom=40, start=80, end=80)
            if is_last:
                set_cell_border(cell, top=none, bottom=line, left=none, right=none)
            else:
                set_cell_border(cell, top=none, bottom=none, left=none, right=none)

    return table


def add_references(doc, refs):
    """参考文献"""
    p = doc.add_heading('参考文献', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_font(run, '黑体', 'Times New Roman', Pt(10.5), bold=True)
    doc.add_paragraph()

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)  # 悬挂缩进
        run = p.add_run(ref)
        set_font(run, '宋体', 'Times New Roman', Pt(10.5))

    doc.add_page_break()


def add_acknowledgement(doc, text):
    """致谢"""
    p = doc.add_heading('谢  辞', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_font(run, '黑体', 'Times New Roman', Pt(16), bold=True)
    doc.add_paragraph()

    p = add_paragraph_with_font(doc, text, '宋体', 'Times New Roman',
                                 Pt(12), indent_first=Cm(0.74), line_spacing=1.5)

    doc.add_page_break()


def add_appendix(doc, title, content):
    """附录
    标题：附录（不加编号，五号、黑体、加粗、居中）
    内容：五号、宋体
    """
    # 附录标题
    p = doc.add_heading(title, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_font(run, '黑体', 'Times New Roman', Pt(10.5), bold=True)
    doc.add_paragraph()

    # 附录内容
    for para in content:
        if isinstance(para, dict) and para.get('type') == 'table':
            add_table_three_line(doc, para['headers'], para['rows'])
        elif isinstance(para, str) and para.startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', para)
            if match:
                caption, path = match.groups()
                add_figure(doc, path, caption, Inches(5))
            else:
                add_body(doc, para)
        else:
            add_body(doc, para)

    doc.add_page_break()


def latex_to_omml(latex_str):
    """将 LaTeX 公式转换为 Word OMML 元素"""
    mathml = latex2mathml.converter.convert(latex_str)
    tree = etree.fromstring(mathml)
    xslt = etree.parse(_MML2OMML_XSL)
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()


def add_formula(doc, formula, formula_num=''):
    """公式
    格式：公式居中，公式编号右对齐
    英文字母和数字为 Times New Roman 体，小四号字
    使用 OMML 原生公式（Word 可编辑）
    """
    # 创建一个表格来实现公式居中、编号右对齐
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置表格无边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # 设置表格宽度为100%
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    # 公式内容（左单元格）- 使用 OMML 原生公式
    cell_formula = table.rows[0].cells[0]
    cell_formula.text = ''
    p = cell_formula.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 尝试将 LaTeX 转换为 OMML
    try:
        omml_elem = latex_to_omml(formula)
        p._element.append(omml_elem)
    except Exception:
        # 转换失败时降级为纯文本
        run = p.add_run(formula)
        set_font(run, '宋体', 'Times New Roman', Pt(12))

    # 公式编号（右单元格）
    cell_num = table.rows[0].cells[1]
    cell_num.text = ''
    p = cell_num.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if formula_num:
        run = p.add_run(formula_num)
        set_font(run, '宋体', 'Times New Roman', Pt(12))

    # 设置单元格边框为无
    for cell in table.rows[0].cells:
        set_cell_border(cell,
                       top={"val": "none", "sz": "0", "space": "0"},
                       bottom={"val": "none", "sz": "0", "space": "0"},
                       left={"val": "none", "sz": "0", "space": "0"},
                       right={"val": "none", "sz": "0", "space": "0"})

    return table


def parse_md_frontmatter(md_text):
    """解析 Markdown 中的 YAML frontmatter"""
    import yaml
    match = re.match(r'^---\s*\n(.*?)\n---\s*\n(.*)', md_text, re.DOTALL)
    if match:
        meta = yaml.safe_load(match.group(1))
        body = match.group(2)
        return meta, body
    return {}, md_text


def parse_md_sections(body):
    """解析 Markdown 正文为 sections 结构"""
    sections = []
    current_section = None
    current_content = []
    in_table = False
    table_lines = []
    in_formula = False
    formula_lines = []

    for line in body.split('\n'):
        stripped = line.strip()

        # 检测公式（$$ 开头和结尾）
        if stripped.startswith('$$'):
            if in_formula:
                # 公式结束
                in_formula = False
                formula_text = '\n'.join(formula_lines)
                # 提取公式编号（如果有）
                formula_num = ''
                num_match = re.search(r'\s*\((\d+[-－]\d+)\)\s*$', formula_text)
                if num_match:
                    formula_num = num_match.group(1)
                    formula_text = formula_text[:num_match.start()].strip()
                current_content.append({'type': 'formula', 'formula': formula_text, 'num': formula_num})
                formula_lines = []
            else:
                # 公式开始
                in_formula = True
                # 检查单行公式 $$ ... $$
                if stripped.endswith('$$') and len(stripped) > 2:
                    formula_text = stripped[2:-2].strip()
                    current_content.append({'type': 'formula', 'formula': formula_text, 'num': ''})
                    in_formula = False
            continue

        if in_formula:
            formula_lines.append(stripped)
            continue

        # 检测表格行
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            table_lines.append(stripped)
            continue

        # 表格结束，解析为结构化数据
        if in_table and not (stripped.startswith('|') and stripped.endswith('|')):
            in_table = False
            if len(table_lines) >= 3:
                # 第一行表头，第二行分隔线，后续为数据行
                headers = [c.strip() for c in table_lines[0].strip('|').split('|')]
                rows = []
                for tl in table_lines[2:]:
                    cells = [c.strip() for c in tl.strip('|').split('|')]
                    rows.append(cells)
                current_content.append({'type': 'table', 'headers': headers, 'rows': rows})
            table_lines = []

        # 检测标题
        heading_match = re.match(r'^(#{1,3})\s+(.*)', line)
        if heading_match:
            if current_section:
                current_section['content'] = current_content
                sections.append(current_section)
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            current_section = {'level': level, 'title': title}
            current_content = []
        elif stripped:
            current_content.append(stripped)

    # 处理末尾表格
    if in_table and len(table_lines) >= 3:
        headers = [c.strip() for c in table_lines[0].strip('|').split('|')]
        rows = []
        for tl in table_lines[2:]:
            cells = [c.strip() for c in tl.strip('|').split('|')]
            rows.append(cells)
        current_content.append({'type': 'table', 'headers': headers, 'rows': rows})

    if current_section:
        current_section['content'] = current_content
        sections.append(current_section)

    return sections


def generate_from_json(content_file, output_file=None):
    """从 JSON 文件生成论文"""
    with open(content_file, 'r', encoding='utf-8') as f:
        content = json.load(f)

    # 默认输出文件名
    if output_file is None:
        output_file = content_file.replace('.json', '.docx')

    doc = create_thesis_template()

    # 封面（无页码）
    cover = content['cover']
    add_cover_page(doc,
        title=cover['title'],
        college=cover['college'],
        major=cover['major'],
        student_id=cover['student_id'],
        name=cover['name'],
        advisor=cover['advisor'],
        date=cover['date'])

    # ===== 分节：摘要部分（罗马数字页码） =====
    section_abstract = add_section_break(doc, 'newPage')
    set_page_number(section_abstract, fmt='roman', start=1)

    # 中文摘要
    cn = content['abstract_cn']
    add_abstract_cn(doc, abstract=cn['content'], keywords=cn['keywords'])

    # 英文摘要
    en = content['abstract_en']
    add_abstract_en(doc, title_en=en['title'], abstract_en=en['content'], keywords_en=en['keywords'])

    # ===== 分节：目录（不编页码） =====
    section_toc = add_section_break(doc, 'newPage')
    add_toc(doc)

    # ===== 分节：正文（阿拉伯数字页码） =====
    section_body = add_section_break(doc, 'newPage')
    set_page_number(section_body, fmt='decimal', start=1)

    # 正文
    for section in content['sections']:
        level = section['level']
        title = section['title']

        if level == 1:
            add_heading_1(doc, title)
        elif level == 2:
            add_heading_2(doc, title)
        elif level == 3:
            add_heading_3(doc, title)

        for para in section.get('content', []):
            if isinstance(para, dict) and para.get('type') == 'table':
                add_table_three_line(doc, para['headers'], para['rows'])
            elif isinstance(para, dict) and para.get('type') == 'figure':
                add_figure(doc, para['path'], para.get('caption', ''), Inches(5))
            elif isinstance(para, dict) and para.get('type') == 'formula':
                add_formula(doc, para['formula'], para.get('num', ''))
            elif isinstance(para, str) and para.startswith('!['):
                # 解析 Markdown 图片语法: ![caption](path)
                match = re.match(r'!\[(.*?)\]\((.*?)\)', para)
                if match:
                    caption, path = match.groups()
                    add_figure(doc, path, caption, Inches(5))
                else:
                    add_body(doc, para)
            else:
                add_body(doc, para)

    # 附录
    for appendix in content.get('appendix', []):
        add_appendix(doc, appendix['title'], appendix.get('content', []))

    # 参考文献
    add_references(doc, content['references'])

    # 致谢
    add_acknowledgement(doc, content['acknowledgement'])

    # 保存
    doc.save(output_file)
    print(f'生成完成: {output_file}')
    return output_file


def generate_from_md(content_file, output_file=None):
    """从 Markdown 文件生成论文"""
    with open(content_file, 'r', encoding='utf-8') as f:
        md_text = f.read()

    meta, body = parse_md_frontmatter(md_text)
    sections = parse_md_sections(body)

    # 默认输出文件名
    if output_file is None:
        output_file = content_file.replace('.md', '.docx')

    doc = create_thesis_template()

    # 封面（无页码）
    add_cover_page(doc,
        title=meta.get('title', ''),
        college=meta.get('college', ''),
        major=meta.get('major', ''),
        student_id=str(meta.get('student_id', '')),
        name=meta.get('name', ''),
        advisor=meta.get('advisor', ''),
        date=meta.get('date', ''))

    # ===== 分节：摘要部分（罗马数字页码） =====
    section_abstract = add_section_break(doc, 'newPage')
    set_page_number(section_abstract, fmt='roman', start=1)

    # 解析中文摘要
    abstract_cn_section = None
    abstract_en_section = None
    references_section = None
    acknowledgement_section = None
    body_sections = []

    appendix_sections = []

    for sec in sections:
        title = sec['title'].strip()
        if title == '摘要':
            abstract_cn_section = sec
        elif title == 'Abstract':
            abstract_en_section = sec
        elif title == '参考文献':
            references_section = sec
        elif title == '致谢':
            acknowledgement_section = sec
        elif title == '附录':
            appendix_sections.append(sec)
        else:
            body_sections.append(sec)

    # 中文摘要
    if abstract_cn_section:
        content = abstract_cn_section['content']
        # 提取关键词
        keywords = ''
        abstract_text = []
        for line in content:
            if line.startswith('**关键词：**') or line.startswith('**关键词:**'):
                keywords = line.replace('**关键词：**', '').replace('**关键词:**', '').strip()
            else:
                abstract_text.append(line)
        add_abstract_cn(doc, abstract='\n'.join(abstract_text), keywords=keywords)

    # 英文摘要
    if abstract_en_section:
        content = abstract_en_section['content']
        keywords_en = ''
        abstract_en_text = []
        for line in content:
            if line.startswith('**Key words:**') or line.startswith('**Key words：**'):
                keywords_en = line.replace('**Key words:**', '').replace('**Key words：**', '').strip()
            else:
                abstract_en_text.append(line)
        add_abstract_en(doc, title_en=meta.get('title', ''),
                       abstract_en='\n'.join(abstract_en_text),
                       keywords_en=keywords_en)

    # ===== 分节：目录（不编页码） =====
    section_toc = add_section_break(doc, 'newPage')
    add_toc(doc)

    # ===== 分节：正文（阿拉伯数字页码） =====
    section_body = add_section_break(doc, 'newPage')
    set_page_number(section_body, fmt='decimal', start=1)

    # 正文
    for section in body_sections:
        level = section['level']
        title = section['title']

        if level == 1:
            add_heading_1(doc, title)
        elif level == 2:
            add_heading_2(doc, title)
        elif level == 3:
            add_heading_3(doc, title)

        for para in section.get('content', []):
            if isinstance(para, dict) and para.get('type') == 'table':
                add_table_three_line(doc, para['headers'], para['rows'])
            elif isinstance(para, dict) and para.get('type') == 'figure':
                add_figure(doc, para['path'], para.get('caption', ''), Inches(5))
            elif isinstance(para, dict) and para.get('type') == 'formula':
                add_formula(doc, para['formula'], para.get('num', ''))
            elif isinstance(para, str) and para.startswith('!['):
                # 解析 Markdown 图片语法: ![caption](path)
                match = re.match(r'!\[(.*?)\]\((.*?)\)', para)
                if match:
                    caption, path = match.groups()
                    add_figure(doc, path, caption, Inches(5))
                else:
                    add_body(doc, para)
            else:
                add_body(doc, para)

    # 附录
    for appendix_sec in appendix_sections:
        add_appendix(doc, appendix_sec['title'], appendix_sec.get('content', []))

    # 参考文献
    if references_section:
        add_references(doc, references_section['content'])

    # 致谢
    if acknowledgement_section:
        add_acknowledgement(doc, '\n'.join(acknowledgement_section['content']))

    # 保存
    doc.save(output_file)
    print(f'生成完成: {output_file}')
    print('提示：在 Word 中右键目录区域，选择"更新域"→"更新整个目录"以生成目录')
    return output_file


if __name__ == '__main__':
    # 命令行参数
    content_file = sys.argv[1] if len(sys.argv) > 1 else 'content.md'
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    # 根据文件扩展名选择函数
    if content_file.endswith('.json'):
        generate_from_json(content_file, output_file)
    else:
        generate_from_md(content_file, output_file)
